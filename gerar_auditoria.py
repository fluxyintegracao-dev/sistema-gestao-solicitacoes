# -*- coding: utf-8 -*-
"""Gera o documento DOCX da Auditoria Técnica, Operacional e Estratégica do Ecossistema Fluxy."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos globais ──────────────────────────────────────────────────────────
styles = doc.styles

# Margens
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# Fonte padrão
style_normal = styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(10)

# Cores
COR_PRIMARIA   = RGBColor(0x00, 0x3A, 0x8C)   # azul escuro
COR_ACENTO     = RGBColor(0xD4, 0x45, 0x00)   # laranja/vermelho para crítico
COR_VERDE      = RGBColor(0x1A, 0x7A, 0x3C)
COR_CINZA      = RGBColor(0x55, 0x55, 0x55)
COR_FUNDO_TBL  = RGBColor(0xE8, 0xEF, 0xF7)   # azul claro para cabeçalho de tabela


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = COR_PRIMARIA
        run.font.size = Pt(22)
        run.bold = True
    return p


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = COR_PRIMARIA
        run.font.size = Pt(14)
        run.bold = True
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = COR_PRIMARIA
        run.font.size = Pt(12)
        run.bold = True
    return p


def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = COR_CINZA
        run.font.size = Pt(11)
        run.bold = True
    return p


def add_paragraph(doc, text, bold=False, italic=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], 'E8EFF7')
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = COR_PRIMARIA

    # Linhas
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            row_cells[ci].text = str(val)
            for para in row_cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 0:
                set_cell_bg(row_cells[ci], 'F7F9FC')

    # Larguras
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_notice(doc, text, tipo='info'):
    """Caixa de destaque visual (simulada com parágrafo indentado e negrito)."""
    cores = {
        'critico': COR_ACENTO,
        'aviso':   RGBColor(0xB5, 0x6F, 0x00),
        'info':    COR_PRIMARIA,
        'ok':      COR_VERDE,
    }
    cor = cores.get(tipo, COR_PRIMARIA)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10)
    run.bold       = True
    run.font.color.rgb = cor
    return p


# ════════════════════════════════════════════════════════════════════════════════
# CAPA
# ════════════════════════════════════════════════════════════════════════════════
add_title(doc, 'AUDITORIA TÉCNICA, OPERACIONAL E ESTRATÉGICA')
add_title(doc, 'ECOSSISTEMA FLUXY')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Análise completa dos repositórios C:\\Fluxy e C:\\Fluxy_Experience')
r.font.color.rgb = COR_CINZA
r.font.size = Pt(11)
r.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Data: 23 de maio de 2026  |  Executor: Claude Sonnet 4.6  |  Branch: dev-v2')
r.font.size = Pt(10)
r.font.color.rgb = COR_CINZA

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# SUMÁRIO EXECUTIVO
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, 'SUMÁRIO EXECUTIVO')
add_paragraph(doc,
    'O ecossistema Fluxy é composto por dois repositórios separados: o ERP interno da construtora '
    '(C:\\Fluxy) e o site comercial imersivo (C:\\Fluxy_Experience). O ERP contém ~82 controllers, '
    '~112 services, ~105 models, 65 migrações JS e 141 páginas frontend. O site Experience '
    'possui viewer 3D com WebXR real, CMS, admin de empreendimentos e integração de leads com o ERP.',
    size=10)
doc.add_paragraph()
add_paragraph(doc,
    'O sistema demonstra maturidade técnica acima da média para desenvolvimento solo, com segurança '
    'implementada de forma real (JWT, MFA TOTP, CSRF, rate limiting, logs de auditoria), arquitetura '
    'de services decompostos por domínio e migrações rastreáveis. Entretanto, apresenta riscos '
    'estruturais críticos que crescem com o sistema: zero testes unitários em ~680 arquivos de código, '
    'dependência total do criador para operação e deploy, ausência de CI/CD, staging e rollback, e '
    'dois sistemas paralelos de gestão de schema no banco de dados.',
    size=10)
doc.add_paragraph()
add_notice(doc,
    'RISCO CRÍTICO: Zero testes unitários + fakeAuth.js referenciado em routes.js + dual schema '
    'management (prepararBanco + migrations). Estes três itens são prioridade máxima antes de '
    'qualquer expansão de funcionalidades.',
    tipo='critico')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 1. VISÃO GERAL DOS PROJETOS
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '1. VISÃO GERAL DOS PROJETOS')

add_h2(doc, '1.1 Repositórios e Stack')
add_table(doc,
    ['Repositório', 'Finalidade', 'Stack', 'Arquitetura'],
    [
        ['C:\\Fluxy', 'ERP interno da construtora', 'Node.js/Express + React/Vite', 'Monolito modular, SPA'],
        ['C:\\Fluxy_Experience', 'Site comercial da construtora', 'Next.js 15 (App Router) + Express API própria', 'SSR/ISR + API REST separada'],
    ],
    col_widths=[3.5, 4.0, 5.5, 4.5]
)

add_h2(doc, '1.2 Métricas de Volume')
add_table(doc,
    ['Artefato', 'C:\\Fluxy (ERP)', 'C:\\Fluxy_Experience'],
    [
        ['Controllers',                   '~82',  '5'],
        ['Services',                      '~112', '4'],
        ['Models (Sequelize)',             '~105', '8'],
        ['Middlewares',                   '13',   '2'],
        ['Migrations',                    '65 JS + 13 SQL', '24'],
        ['Páginas frontend',              '141',  '~15 rotas Next.js'],
        ['Componentes reutilizáveis',     '~12',  '~40'],
        ['Arquivos JS/TS totais (src)',   '404',  '78 TSX + 31 JS (API)'],
        ['Testes unitários',              '0',    '0'],
        ['Testes E2E (Playwright)',       '6 specs', '0'],
        ['Scripts de validação manual',  '6',    '0'],
    ],
    col_widths=[5.5, 5.0, 5.0]
)

add_h2(doc, '1.3 Organização Estrutural')
add_h3(doc, 'Pontos positivos')
for item in [
    'Separação clara entre os dois produtos com APIs independentes.',
    'Backend do ERP com organização por camadas (controllers / services / models / middlewares) consistente.',
    'Nomenclatura de arquivos padronizada e descritiva em todo o projeto.',
    'Fluxy_Experience com arquitetura Next.js 15 moderna (App Router, TypeScript, RSC).',
    'Experience_API separada do ERP — separação correta de contextos e segurança.',
    'Módulo Fiscal isolado em backend/src/modules/fiscal/ com estrutura completa própria.',
]:
    add_bullet(doc, item)

add_h3(doc, 'Problemas estruturais')
for item in [
    'Roteamento via arquivo monolítico routes.js (~700+ linhas) em vez de arquivos por domínio — cresce indefinidamente.',
    '141 páginas frontend com apenas ~12 componentes reutilizáveis — baixíssima componentização.',
    'Ausência de CI/CD: nenhum pipeline automatizado (.github/, .gitlab-ci.yml) identificado.',
    'docker-compose.yml cobre apenas o MySQL local — não cobre backend, frontend ou Redis.',
    'Nenhum .nvmrc ou campo engines no package.json — versão Node (24.11.1) não está declarada no repositório.',
    'Ausência de README.md na raiz e de CLAUDE.md — onboarding depende de AGENTS.md.',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 2. ANÁLISE DE MATURIDADE POR MÓDULO
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '2. ANÁLISE DE MATURIDADE POR MÓDULO')

add_h2(doc, '2.1 ERP — C:\\Fluxy')
add_table(doc,
    ['Módulo', 'Classificação Real', 'Evidência / Observação'],
    [
        ['Solicitações',      'Produção Estável',    'Controllers + 5 services decompostos, SLA configurável, automação de status, relatório operacional, integração com provisionamento (commit 6035ac1)'],
        ['Compras / Cotações','Produção Parcial',    'PDF via Puppeteer, cotação de fornecedor, pedido de compra, relatórios implementados. Cotação avulsa nullable indica refinamento em andamento'],
        ['Financeiro (Títulos)','Produção Parcial',  '20+ controllers: títulos, cartão, fatura, caixa, transferência, conciliação, DRE, centros de custo. Integrações bancárias ainda em sandbox'],
        ['DRE',               'MVP',                 '10 endpoints em RelatorioFinanceiroController, relatorioFinanceiroService com gerarDreGerencial, comparativo mensal e diagnóstico. Funcional mas não auditado operacionalmente'],
        ['Intercompany',      'Estrutura Inicial',   'Migrações 202605210002-3 criadas (mai/2026). Colunas adicionadas a titulos_financeiros e movimentos_financeiros. Sem uso operacional validado'],
        ['Provisionamento',   'Produção Parcial',    '3 controllers, services com lifecycle de status, dashboards, categorias macro, histórico e permissões. Integração com solicitações recente'],
        ['RH/DP',             'MVP',                 '6 controllers (colaboradores, documentos, importações, apurações, fechamentos, empresas), services completos. Operação não confirmada externamente'],
        ['Obras',             'Produção Parcial',    'ObraController + obraGestaoService, centro de custo, apropriação, resultado de obras'],
        ['Comercial',         'MVP',                 'Empreendimentos, unidades, contratos, tabelas de preço, D4Sign, compradores, testemunhas. Módulo jovem (migrations abr/mai 2026)'],
        ['CRM',               'Experimental',        '10 controllers, 9 services, webhooks Meta+Google, automações em runtime, 4 fases entregues. Módulo muito amplo para adoção validada'],
        ['Fiscal',            'MVP Desabilitado',    'FISCAL_MODULE_ENABLED=false. Módulo completo em modules/fiscal/: 7 controllers, ~15 services, 10 models, SEFAZ DFe, DANFE, criptografia'],
        ['CNAB / Boletos Caixa','Experimental',      'Geração de remessa CNAB 240 e importação de retorno implementadas. CAIXA_BOLETO_AMBIENTE=TESTE, CAIXA_BOLETO_HOMOLOGADO=false'],
        ['Pagamentos BB',     'Experimental',        'OAuth2 + mTLS, payload mapper, approval flow com MFA step-up. BB_PAYMENTS_ENV=sandbox. Não em produção'],
        ['Sienge Integration','MVP',                 'Config, fila, log, mapeamento implementados. REST API com Basic/token auth. CSV de carga inicial. Operacional por configuração'],
        ['Comunicação Interna','MVP',                'Chat interno, participantes, anexos, reply, live updates via SSE (EventSource). Estruturalmente completo'],
        ['MFA / Segurança',   'Produção Parcial',    'Ver seção 4'],
        ['Analytics / Ops',   'Estrutura Inicial',   'OPS_ENABLED=false. Heartbeat e métricas existem mas desabilitados. Sem observabilidade ativa'],
        ['Mobile (Capacitor)','Experimental',        'capacitor.config.json + runtime.js + dependências v8 instaladas. Projetos nativos (android/ ios/) NÃO gerados. Não publicado em stores'],
    ],
    col_widths=[3.8, 3.8, 9.9]
)

add_h2(doc, '2.2 Fluxy Experience — C:\\Fluxy_Experience')
add_table(doc,
    ['Módulo', 'Classificação Real', 'Evidência / Observação'],
    [
        ['Viewer 3D',           'Produção Parcial',  'React Three Fiber + Three.js real. Orbit + first-person com física de colisão. Hotspots 3D. Fallback procedural. Performance adaptada para touch. Analytics integrado'],
        ['WebXR (VR imersivo)', 'Experimental',      'xr.requestSession("immersive-vr") implementado, renderer com gl.xr.enabled=true. Código real, funcional em hardware compatível. Não testado em produção com headset'],
        ['CMS / Site Editor',  'MVP',                'SiteEditorController, versionamento (migration 024), sections por página, drafts, publish/snapshot. Admin funcional'],
        ['SEO',                 'Implementado',      'robots.ts, sitemap.ts, SchemaOrg.tsx. ISR do Next.js garante indexação'],
        ['Analytics',          'Produção Parcial',   'AnalyticsController, ExperiencePageEvent, Google Analytics, GTM, Meta Pixel, tracking granular de hotspots 3D, VR, lead gate, UTMs'],
        ['Leads / Formulário', 'Produção Estável',   'ExperienceLead com sync para ERP core via leadCoreSyncService. Contexto de hotspot/unidade/simulação. Dedup via 409. Status: pending/synced/failed/duplicate'],
        ['Simulador Financiamento','MVP',            'SimuladorFinanciamento.tsx + SimuladorLeadGate.tsx presentes. Lógica de cálculo não auditada em profundidade'],
        ['Blog',               'MVP',                'Model, admin CRUD, páginas públicas implementados'],
        ['Admin',              'MVP',                'AdminShell, CRUD empreendimentos, upload imagens/modelos 3D, editor de hotspots, editor de transform 3D, checklist de publicação'],
        ['Unidades / Mapa',    'Estrutura Inicial',  'ExperienceUnidade, MapaUnidades.tsx, migration 003 presente. Profundidade da interface a validar'],
        ['Sync com ERP Core',  'Produção Parcial',   'syncService + leadCoreSyncService + experienceSyncRouter no ERP. Lead sync e empreendimento sync implementados'],
    ],
    col_widths=[3.8, 3.8, 9.9]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 3. ANÁLISE DE ARQUITETURA
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '3. ANÁLISE DE ARQUITETURA')

add_h2(doc, '3.1 Pontos Arquiteturais Positivos')
for item in [
    'Services de solicitações decompostos por responsabilidade (criarSolicitacao.js, listarSolicitacoes.js, atualizarStatus.js como arquivos separados) — evita god service.',
    'Modelo de autorização granular: middleware auth.js carrega permissões financeiras, RH/DP e áreas de permissão em um único Promise.all por requisição.',
    'Experience com API própria: banco separado, sincronização controlada, sem exposição direta do banco do ERP ao site público.',
    'Rate limiting com janelas configuráveis por contexto e registro de eventos de segurança.',
    'Módulo Fiscal autocontido em modules/fiscal/ com controllers, services, models e routes próprios.',
]:
    add_bullet(doc, item)

add_h2(doc, '3.2 Gargalos e Pontos Frágeis')
add_table(doc,
    ['Problema', 'Severidade', 'Impacto'],
    [
        ['routes.js monolítico (~700+ linhas, 80+ módulos registrados)',         'Alta',   'Ponto de acoplamento que cresce indefinidamente; impossível testar isoladamente'],
        ['DUAL SCHEMA MANAGEMENT: prepararBanco() em app.js + runMigrations.js', 'Crítica','Estado do banco depende de dois sistemas independentes. Inconsistência silenciosa em falhas parciais'],
        ['141 páginas com ~12 componentes reutilizáveis',                        'Alta',   'Mudanças de UI requerem edição em dezenas de arquivos. Custo de manutenção cresce linearmente'],
        ['Ausência de camada de repositório',                                    'Média',  'Controllers/services acessam Sequelize diretamente. Acopla lógica de domínio ao ORM'],
        ['CRM com 10 controllers sem estrutura de domínio clara',                'Média',  'Acoplamento crescente; difícil entender fluxos de automação complexos'],
        ['Ausência de tratamento centralizado de erros',                         'Média',  'Respostas de erro inconsistentes entre módulos'],
        ['fakeAuth.js referenciado em routes.js',                                'Crítica','Se ativável em produção, elimina toda autenticação'],
    ],
    col_widths=[6.5, 2.5, 8.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 4. ANÁLISE DE SEGURANÇA REAL
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '4. ANÁLISE DE SEGURANÇA REAL')

add_h2(doc, '4.1 Segurança Implementada')
add_table(doc,
    ['Mecanismo', 'Status Real', 'Detalhe'],
    [
        ['JWT',                'Implementado',   'jsonwebtoken, dual-mode: Bearer header OU cookie HttpOnly. Exclui senha, mfa_totp_secret da query'],
        ['Cookies HttpOnly/Secure', 'Implementado', 'AUTH_COOKIE_SECURE=true, AUTH_COOKIE_SAME_SITE=lax. Cookie configurável por domínio'],
        ['MFA (TOTP)',         'Implementado',   'otplib, SHA1, 6 dígitos, 30s period, 1 period tolerance. QR code via qrcode lib. requireMfaCompletion bloqueia rotas'],
        ['CSRF',              'Implementado',   'Double-submit: cookie fluxy_csrf vs header X-CSRF-Token. Aplicado apenas em auth_mode=cookie'],
        ['Rate Limiting',     'Implementado',   'Custom com Redis (fallback in-memory). Login: 5/15min, upload: 30/5min, crítico: 60/1min, senha: 5/15min. Logs via SecurityEventLog'],
        ['Helmet / CSP',      'Implementado',   'CSP: defaultSrc self, frameAncestors none, objectSrc none. HSTS com preload em produção. noSniff, referrerPolicy same-origin, frameguard deny'],
        ['CORS dinâmico',     'Implementado',   'Allowlist carregada do banco (config de instalação), não de env estático. Suporte a wildcards. Bloqueios logados'],
        ['Upload validation', 'Implementado',   'Validação binária de MIME type real (assertFileBinaryMatchesProfile). Extensões perigosas (.html, .js, .svg) servidas como attachment com sandbox CSP'],
        ['Logs de auditoria', 'Implementado',   'SecurityEventLog model. Eventos: AUTH_TOKEN_MISSING, AUTH_TOKEN_INVALID, RATE_LIMIT_BLOCKED, CSRF_MISMATCH registrados com IP, user, recurso'],
        ['Permissões por área','Implementado',   'PermissoesAreas, SetorPermissao, moduleAccess, resourceAccess, crmAccess — camadas múltiplas de controle'],
        ['ClamAV',            'Parcial',        'Código implementado (clamavService.js). CLAMAV_ENABLED=false por padrão — NÃO ativo em produção'],
        ['Refresh token',     'Ausente',        'JWT de 8h sem silent refresh. Expiração força re-login manual'],
        ['IP allowlist',      'Ausente',        'Sem restrição por IP em nível de aplicação'],
    ],
    col_widths=[3.5, 3.0, 11.0]
)

add_h2(doc, '4.2 Riscos de Segurança')
add_table(doc,
    ['Risco', 'Nível', 'Ação Necessária'],
    [
        ['fakeAuth.js referenciado em routes.js',          'CRÍTICO', 'Investigar condições de ativação. Se ativável em produção via env/flag, remover imediatamente'],
        ['S3 CORS AllowedOrigins: ["*"]',                  'MÉDIO-ALTO', 'Restringir para as origens reais da aplicação (jrfluxy.com.br, Vercel domain)'],
        ['Redis REDIS_REQUIRED=false',                     'MÉDIO',   'Rate limiting cai para in-memory sem aviso em restart. Tornar Redis obrigatório em produção'],
        ['ClamAV desabilitado',                            'MÉDIO',   'Uploads sem scan de vírus em produção. Avaliar habilitação'],
        ['CORS_ALLOWED_ORIGINS vazio em .env.example',     'MÉDIO',   'Se mal configurado em produção, pode aceitar qualquer origem'],
        ['Boletos Caixa em TESTE/não homologado',          'BAIXO',   'Nenhum risco atual; formalizar processo de ativação antes de habilitar'],
        ['BB Payments em sandbox',                         'BAIXO',   'Mesma situação — processo formal antes de habilitar'],
        ['docker-compose expõe 0.0.0.0:3306',             'BAIXO',   'Risco em VMs compartilhadas ou desenvolvimento com rede exposta'],
    ],
    col_widths=[6.0, 2.5, 9.0]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 5. INFRAESTRUTURA E DEVOPS
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '5. INFRAESTRUTURA E DEVOPS')

add_table(doc,
    ['Componente', 'Status', 'Observação'],
    [
        ['MySQL (AWS RDS)',      'Ativo em produção', 'DB_HOST configurável no env. Sem configuração de pool explícita no Sequelize'],
        ['AWS S3',              'Ativo em produção', 'AWS SDK v3. CORS atualmente AllowedOrigins: ["*"] — requer correção'],
        ['Redis',               'Opcional',          'REDIS_REQUIRED=false. Fallback para in-memory Map em caso de indisponibilidade'],
        ['PM2',                 'Ativo',             'pm2 start backend-solicitacoes no EC2. Scripts pm2:start/restart/logs em package.json da Experience API'],
        ['Nginx',               'Ativo (inferido)',  'Proxy para 127.0.0.1:8000 no EC2. Não versionado no repositório'],
        ['Vercel',              'Ativo',             'Frontend ERP (Vite build) + Fluxy_Experience (Next.js). Auto-deploy via git push'],
        ['CI/CD',               'AUSENTE',           'Nenhum pipeline automatizado. Deploy manual: git pull + npm install + pm2 restart'],
        ['Dockerfile',          'AUSENTE',           'Nenhum Dockerfile para backend ou frontend. docker-compose apenas para MySQL local'],
        ['Staging',             'AUSENTE',           'Sem ambiente de homologação. Mudanças vão direto para produção'],
        ['Rollback',            'Manual',            'Sem automação. Requer git revert + pm2 restart manual'],
        ['Backups',             'Não documentado',   'Não identificado no repositório. Crítico para dados financeiros'],
        ['Observabilidade',     'Desabilitada',      'OPS_ENABLED=false. Heartbeat/métricas existem no código mas não ativos'],
        ['Node.js versão',      'Não declarada',     'Node 24.11.1 em produção (AGENTS.md), mas sem .nvmrc ou engines no package.json'],
    ],
    col_widths=[3.5, 3.0, 11.0]
)

add_notice(doc,
    'A infraestrutura é adequada para operação interna atual, mas é completamente dependente de configuração manual. '
    'Um deploy errado, sem staging e sem rollback automatizado, pode parar o ERP da empresa sem caminho rápido de recuperação.',
    tipo='aviso')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 6. DÍVIDA TÉCNICA
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '6. DÍVIDA TÉCNICA')

add_h2(doc, '6.1 Dívidas Críticas — Resolver Imediatamente')
add_table(doc,
    ['Dívida', 'Severidade', 'Impacto', 'Urgência'],
    [
        ['Zero testes unitários (~680 arquivos de lógica)',          'Crítica', 'Qualquer mudança pode quebrar fluxos em produção silenciosamente. Superfície enorme de regressão',     'Imediata'],
        ['fakeAuth.js referenciado em routes.js',                    'Crítica', 'Potencial bypass total de autenticação se ativável em produção',                                       'Imediata'],
        ['Dual schema management (prepararBanco + runMigrations)',   'Crítica', 'Estado do banco depende de dois sistemas independentes. Falha parcial gera inconsistência silenciosa', 'Imediata'],
        ['Ausência de CI/CD',                                        'Alta',    'Deploys manuais sem gate de qualidade. Risco humano crescente',                                        'Alta'],
    ],
    col_widths=[5.5, 2.5, 6.5, 2.5]
)

add_h2(doc, '6.2 Dívidas Médias — Resolver em 3 a 6 Meses')
add_table(doc,
    ['Dívida', 'Impacto'],
    [
        ['141 páginas / ~12 componentes reutilizáveis',   'Custo de manutenção alto. Mudanças de UI exigem edição em dezenas de arquivos'],
        ['CRM com 10 controllers sem domínio claro',      'Acoplamento crescente; fluxos de automação difíceis de rastrear e debugar'],
        ['Redis não obrigatório',                         'Rate limiting perde persistência em restart. Contornável em cenários de carga'],
        ['Ausência de error handler centralizado',        'Respostas de erro inconsistentes entre módulos; dificulta diagnóstico de produção'],
        ['Ausência de staging',                           'Impossível validar mudanças antes de produção'],
        ['S3 CORS wildcard',                              'Superfície desnecessariamente ampla'],
        ['Sem .nvmrc / engines declaration',              'Versão Node divergente entre desenvolvedores e produção'],
    ],
    col_widths=[5.5, 12.0]
)

add_h2(doc, '6.3 Dívidas Baixas — Podem Aguardar')
for item in [
    'Baixa componentização do frontend — produtividade de desenvolvimento',
    'Ausência de documentação de API (Swagger/Postman) — onboarding lento de novos devs',
    'DEPLOY_QUESTIONS.md desatualizado — não cobre estado atual de migrações, S3 e segurança',
    'Módulo Fiscal desabilitado ocupando espaço cognitivo sem gerar valor operacional',
    'docker-compose.yml incompleto — não cobre stack completa de desenvolvimento',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 7. DEPENDÊNCIA DO CRIADOR
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '7. DEPENDÊNCIA DO CRIADOR')

add_h2(doc, '7.1 O que está documentado')
for item in [
    '.env.example detalhado e bem comentado para ambos os projetos.',
    'Nomes de arquivos e funções são descritivos — o código é legível.',
    'Diretório docs/ com ~50 arquivos: arquitetura, segurança, módulos, regras de negócio, contexto, logs de desenvolvimento.',
    'AGENTS.md funciona como documento de onboarding principal.',
    'Commits com mensagens descritivas e rastreáveis.',
]:
    add_bullet(doc, item)

add_h2(doc, '7.2 O que não está documentado')
for item in [
    'Regras de negócio implícitas (por que determinados setores têm comportamentos especiais, quais SLAs se aplicam a cada fluxo, por que o Sienge tem configuração de auto_vincular_credor).',
    'Processo de deploy — não existe runbook automatizado ou script documentado passo a passo.',
    'Processo de backup e recuperação de desastre.',
    'Ordem de execução e dependências entre módulos para onboarding de novos desenvolvedores.',
    'Quais módulos estão realmente em uso operacional versus experimentais — a documentação de planos de módulos não equivale a documentação de uso real.',
    'Comportamentos corretos esperados (zero testes = zero contratos de comportamento).',
    'Documentação de API (sem Swagger/Postman/OpenAPI).',
]:
    add_bullet(doc, item)

add_h2(doc, '7.3 Estimativa de Onboarding')
add_table(doc,
    ['Perfil', 'Tempo estimado para produtividade', 'Principais barreiras'],
    [
        ['Dev sênior (Node.js/React)',  '2 a 4 semanas', 'Entender regras de negócio implícitas, dual schema management, roteamento monolítico'],
        ['Dev pleno',                  '6 a 8 semanas', 'Acima + entender modelo de permissões, módulos financeiros e integrações bancárias'],
        ['Dev júnior',                 '3+ meses',      'Risco de nunca ter confiança para fazer mudanças em módulos financeiros sem supervisão'],
    ],
    col_widths=[4.5, 5.0, 8.0]
)

add_h2(doc, '7.4 Nível Atual de Institucionalização')
add_paragraph(doc,
    'O sistema está no meio do caminho entre "projeto pessoal avançado" e "produto institucional". '
    'A infraestrutura de segurança e a organização do código indicam intenção institucional. '
    'A ausência de testes, CI/CD, documentação operacional e runbook de deploy indica que a '
    'velocidade de construção superou a capacidade de institucionalizar. O sistema é tecnicamente '
    'bem construído em muitas partes, mas organicamente dependente do criador para operação, '
    'deploy, decisão arquitetural e conhecimento de regras de negócio.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 8. ANÁLISE DO FLUXY EXPERIENCE
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '8. ANÁLISE DO FLUXY EXPERIENCE')

add_h2(doc, '8.1 O que realmente está funcionando')

add_h3(doc, 'Viewer 3D — a peça técnica mais madura do Experience')
for item in [
    'Carregamento de modelos GLB/GLTF reais via useGLTF (React Three Fiber).',
    'Modo orbit (rotação/zoom) e modo first-person com física de colisão e limites configuráveis por empreendimento.',
    'Hotspots 3D com posicionamento espacial real e teleporte entre pontos.',
    'Fallback procedural: edifício gerado matematicamente quando não há modelo GLB cadastrado.',
    'Performance adaptada para touch (shadows desligadas, DPR reduzido, controles por toque).',
    'Analytics granular integrado: hotspot_click, first_person_start, first_person_teleport, vr_session_start.',
]:
    add_bullet(doc, item)

add_h3(doc, 'WebXR — código real, não marketing')
for item in [
    'xr.requestSession("immersive-vr") implementado com optionalFeatures: ["local-floor", "bounded-floor"].',
    'renderer.xr.enabled = true configurado no Canvas.',
    'Detecção de suporte via navigator.xr.isSessionSupported antes de exibir botão.',
    'Funcional em hardware compatível (headsets WebXR). Não validado em produção com dispositivo físico nesta auditoria.',
]:
    add_bullet(doc, item)

add_h3(doc, 'Integração com ERP Core')
for item in [
    'leadCoreSyncService faz POST autenticado com X-Experience-Lead-Secret para o ERP.',
    'Trata duplicatas via 409 com duplicateId.',
    'Persiste core_sync_status com estados explícitos: pending/synced/failed/duplicate/disabled.',
    'Sincronização de empreendimentos via experienceSyncRouter com X-Experience-Sync-Key.',
]:
    add_bullet(doc, item)

add_h2(doc, '8.2 O que ainda é parcial ou experimental')
for item in [
    'Simulador de financiamento: componentes existem (SimuladorFinanciamento.tsx + SimuladorLeadGate.tsx). Lógica de cálculo de parcelas não auditada em profundidade.',
    'Mapa de unidades: componente e model existem, mas profundidade da interface de edição e interatividade pública não verificada.',
    'VR imersivo com headset físico: código correto mas dependência de hardware não testado nesta auditoria.',
    'CMS avançado: versionamento e drafts implementados, mas fluxo de publicação para seções complexas a validar.',
]:
    add_bullet(doc, item)

add_h2(doc, '8.3 O que é visão futura não materializada')
for item in [
    'AR (Augmented Reality) — sem código identificado.',
    'Tour imersivo multi-ambiente real além dos hotspots existentes.',
    'Integração com dados de obra em tempo real (andamento de obra, câmeras, etc.).',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 9. ANÁLISE ESTRATÉGICA
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '9. ANÁLISE ESTRATÉGICA')

add_h2(doc, '9.1 O que o Fluxy realmente é hoje')
add_paragraph(doc,
    'O Fluxy hoje é um ERP operacional interno customizado para a operação específica de uma '
    'construtora, com elementos embrionários de produto comercial que ainda não chegaram a MVP '
    'comercializável.')
doc.add_paragraph()

add_table(doc,
    ['Afirmação', 'Verdadeiro?', 'Justificativa'],
    [
        ['ERP operacional interno',        'SIM',          'Módulos de solicitações, compras, financeiro e obras com evidência de uso real'],
        ['Produto comercial pronto',       'NÃO',          'Sem multi-tenancy, sem billing, sem onboarding automatizado, sem SLA documentado'],
        ['SaaS pronto para venda',         'NÃO',          'Intensamente moldado para a operação específica da construtora'],
        ['Base técnica para produto',      'SIM (potencial)', 'Arquitetura, segurança e estrutura de dados são sólidas como ponto de partida'],
        ['Laboratório tecnológico',        'PARCIALMENTE', 'Fiscal, CNAB, BB Payments, CRM automações — experimentos reais com código funcional'],
        ['Sistema dependente do criador',  'SIM',          'Operação, deploy, regras de negócio — tudo centralizado em uma pessoa'],
    ],
    col_widths=[5.0, 2.5, 10.0]
)

add_h2(doc, '9.2 Nível Real de Prontidão Comercial')
add_paragraph(doc,
    'Baixo — não por falta de qualidade técnica, mas por:')
for item in [
    'Ausência de isolamento multi-tenant (dados de um cliente ficam no mesmo banco que os de outro).',
    'Configurações hardcoded para a operação desta construtora específica.',
    'Ausência de documentação para usuário/cliente final.',
    'Onboarding dependente do criador — nenhum processo automatizado de instalação.',
    'Módulos críticos (CNAB, BB Payments, Fiscal) ainda não em produção real.',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 10. O QUE DEVERIA PARAR DE SER EXPANDIDO AGORA
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '10. O QUE DEVERIA PARAR DE SER EXPANDIDO AGORA')

add_table(doc,
    ['Área', 'Justificativa para pausar'],
    [
        ['Módulo Fiscal',                    'Desabilitado explicitamente, código existente e complexo. Adicionar mais sem ativar aumenta dívida sem retorno operacional'],
        ['CRM — novas automações/integrações','10 controllers, automações de IA, webhooks Meta+Google — módulo jovem sem adoção operacional validada. Estabilizar antes de expandir'],
        ['Novas integrações bancárias',      'Antes de homologar Caixa e BB em produção, não expandir para outros bancos ou modalidades'],
        ['Fluxy Experience — novas features','Analytics, simulador, CMS e 3D estão em boa forma. Consolidar o que existe antes de adicionar'],
        ['Mobile (stores)',                  'Sem projetos nativos gerados. Não iniciar publicação antes de estabilizar o core e ter testes'],
        ['Qualquer novo módulo de ERP',      'A relação features/testes está perigosamente desequilibrada. Cada novo módulo aumenta a superfície sem proteção de testes'],
    ],
    col_widths=[5.0, 12.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 11. ROADMAP — PRÓXIMOS 12 MESES
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '11. ROADMAP — PRÓXIMOS 12 MESES')
add_paragraph(doc,
    'Roadmap baseado em estabilidade operacional e governança — não em expansão de funcionalidades.')

add_h2(doc, 'Q3 2026 — Fundação (meses 1 a 3)')
add_table(doc,
    ['#', 'Ação', 'Risco de não executar'],
    [
        ['1', 'Investigar e resolver fakeAuth.js em routes.js',           'Possível bypass de autenticação em produção'],
        ['2', 'Unificar dual schema management (eliminar prepararBanco)', 'Inconsistência de schema em incidentes de restart parcial'],
        ['3', 'Testes unitários nos services críticos (Solicitações, Compras, Financeiro) — meta: 40% de cobertura', 'Qualquer mudança pode quebrar produção silenciosamente'],
        ['4', 'CI/CD básico: GitHub Actions com lint + testes + deploy',  'Deploys manuais com risco humano crescente'],
        ['5', 'Tornar Redis obrigatório em produção',                     'Rate limiting sem persistência entre restarts'],
    ],
    col_widths=[0.7, 8.5, 8.3]
)

add_h2(doc, 'Q4 2026 — Governança e Segurança (meses 4 a 6)')
add_table(doc,
    ['#', 'Ação', 'Risco de não executar'],
    [
        ['6', 'Corrigir S3 CORS AllowedOrigins para lista explícita', 'Superfície desnecessária, risco em pré-signed URLs vazadas'],
        ['7', 'Backups automatizados com teste de restore trimestral', 'Dados financeiros sem recuperação confirmada em caso de falha'],
        ['8', 'Staging environment (réplica de produção)',            'Impossível validar mudanças antes de produção'],
        ['9', 'Error handler centralizado no backend',               'Erros não capturados, respostas inconsistentes'],
        ['10','Runbook de produção documentado (deploy, backup, rollback)', 'Só o criador sabe operar o sistema'],
    ],
    col_widths=[0.7, 8.5, 8.3]
)

add_h2(doc, 'Q1 2027 — Estabilização (meses 7 a 9)')
add_table(doc,
    ['#', 'Ação', 'Risco de não executar'],
    [
        ['11', 'Componentização do frontend (extrair componentes das 30 páginas mais usadas)', 'Custo de manutenção UI crescendo linearmente'],
        ['12', 'Homologação formal Caixa Econômica (CNAB/boletos)',                           'Funcionalidade bloqueada sem processo formal'],
        ['13', 'Documentação de API (Swagger ou Postman collection)',                         'Onboarding de segundo desenvolvedor demorado'],
        ['14', 'Modularizar roteamento (extrair rotas do routes.js por domínio)',             'Arquivo monolítico atingindo limite de manutenibilidade'],
        ['15', 'Declarar versão Node em .nvmrc e engines',                                   'Divergência silenciosa entre ambientes'],
    ],
    col_widths=[0.7, 8.5, 8.3]
)

add_h2(doc, 'Q2 2027 — Equipe e Produto (meses 10 a 12)')
add_table(doc,
    ['#', 'Ação', 'Justificativa'],
    [
        ['16', 'Contratar segundo desenvolvedor (sênior, Node.js/React, cultura de testes)', 'Sem isso, as prioridades anteriores se arrastam indefinidamente'],
        ['17', 'Avaliar adoção operacional do CRM (métricas de uso)',                        'Medir antes de continuar investindo no módulo mais complexo'],
        ['18', 'Planejamento de multi-tenancy se houver intenção comercial',                 'Quanto mais o sistema cresce acoplado, mais cara a migração futura'],
    ],
    col_widths=[0.7, 8.5, 8.3]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 12. PAPEL FUTURO DO FUNDADOR TÉCNICO
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '12. PAPEL FUTURO DO FUNDADOR TÉCNICO')

add_h2(doc, '12.1 Diagnóstico atual')
add_paragraph(doc,
    'O fundador técnico está hoje acumulando os papéis de desenvolvedor, arquiteto, DevOps, product manager e suporte '
    'simultaneamente. Isso é sustentável em fase inicial mas é o principal fator de risco de continuidade operacional do sistema.')

add_h2(doc, '12.2 Transição recomendada')
add_table(doc,
    ['Horizonte', 'Papel', 'Foco'],
    [
        ['Curto prazo (0–6 meses)',  'Desenvolvedor + arquiteto',        'Continuar no desenvolvimento, mas dedicar 20–30% do tempo a testes, documentação e CI/CD'],
        ['Médio prazo (6–18 meses)', 'Arquiteto + product owner técnico', 'Responsável por decisões arquiteturais, não por toda a implementação. Contratar e transferir execução'],
        ['Longo prazo (18m+)',       'CTO interno / tech lead',           'Governança, priorização de produto, inovação estratégica — não execução diária de código'],
    ],
    col_widths=[4.0, 5.0, 8.5]
)

add_h2(doc, '12.3 Perfil ideal para próxima contratação')
for item in [
    'Desenvolvedor full-stack sênior com experiência em Node.js/Express + React.',
    'Familiaridade com domínio ERP/financeiro (fluxos de aprovação, contabilidade básica).',
    'Cultura de testes — este é o requisito diferenciador. Um dev sênior sem cultura de testes piora a dívida existente.',
    'Capaz de absorver documentação técnica e operar independentemente após 4 a 6 semanas de onboarding.',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 13. CONCLUSÃO FINAL
# ════════════════════════════════════════════════════════════════════════════════
add_h1(doc, '13. CONCLUSÃO FINAL')

add_h2(doc, '13.1 Estado Real Atual')
add_table(doc,
    ['Dimensão', 'Avaliação'],
    [
        ['Maturidade técnica geral',    'Acima da média para desenvolvimento solo. Segurança, arquitetura de services e organização estrutural são pontos fortes reais'],
        ['Cobertura de testes',         'Crítica. Zero testes unitários em ~680 arquivos. 6 E2E specs e 6 scripts de validação manuais existem mas são insuficientes'],
        ['Dependência do criador',      'Alta. Operação, deploy, regras de negócio e knowledge — tudo centralizado. Principal risco de continuidade'],
        ['Segurança implementada',      'Boa. JWT, MFA TOTP, CSRF, rate limiting, logs de auditoria — implementados de verdade, não apenas configurados'],
        ['Infraestrutura operacional',  'Funcional mas frágil. Manual, sem staging, sem rollback, sem observabilidade ativa'],
        ['Módulos prontos para produção','Core operacional: Solicitações, Compras, Financeiro base, Obras. Restante em graus variados de maturidade'],
        ['Prontidão comercial',         'Baixa. ERP customizado para operação específica. Requer 12–18 meses de generalização antes de comercializar'],
        ['Sustentabilidade',            'Adequada no curto prazo; frágil no médio prazo sem testes, CI/CD e segundo desenvolvedor'],
        ['Institucionalização',         'Parcial. Código bem estruturado, mas processo operacional dependente do criador'],
    ],
    col_widths=[4.5, 13.0]
)

add_h2(doc, '13.2 Síntese dos Riscos por Prioridade')
add_table(doc,
    ['Risco', 'Nível', 'Consequência se não resolvido'],
    [
        ['Zero testes unitários',                          'CRÍTICO', 'Regressões silenciosas em produção financeira, sem capacidade de detectar antes do impacto ao usuário'],
        ['fakeAuth.js em routes.js',                       'CRÍTICO', 'Possível bypass total de autenticação'],
        ['Dual schema management',                         'CRÍTICO', 'Inconsistência de banco de dados em incidentes, difícil de diagnosticar'],
        ['Dependência total do criador para operação',     'ALTO',    'Indisponibilidade do criador = sistema sem suporte nem evolução'],
        ['Ausência de CI/CD e staging',                   'ALTO',    'Um deploy errado para o ERP da empresa sem caminho de rollback rápido'],
        ['Backups não documentados',                      'ALTO',    'Dados financeiros sem recuperação confirmada em caso de falha de infraestrutura'],
    ],
    col_widths=[6.0, 2.5, 9.0]
)

add_h2(doc, '13.3 Potencial Estratégico Real')
add_paragraph(doc,
    'O potencial estratégico é real e está fundamentado em:')
for item in [
    'Base técnica sólida: o código do core é bem estruturado e pode escalar.',
    'Viewer 3D com WebXR é um diferencial genuíno para o mercado imobiliário — poucos sistemas do segmento têm isso.',
    'Cobertura funcional relevante: o ERP cobre o núcleo operacional de uma construtora (solicitações, compras, financeiro, obras).',
    'Integração Experience + ERP com sync de leads e empreendimentos é um diferencial de produto para vendas imobiliárias.',
]:
    add_bullet(doc, item)
doc.add_paragraph()
add_paragraph(doc,
    'Mas transformar esse potencial em produto comercial ou em sistema institucional sustentável '
    'requer a decisão estratégica central dos próximos 12 meses:')
doc.add_paragraph()
add_notice(doc,
    'Escolher entre continuar expandindo features (risco crescente, retorno de curto prazo) '
    'ou parar de crescer para fortalecer a fundação (custo de curto prazo, sustentabilidade de longo prazo). '
    'As duas coisas em paralelo, com uma pessoa só, não são viáveis.',
    tipo='critico')

# ════════════════════════════════════════════════════════════════════════════════
# RODAPÉ
# ════════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Fluxy — Auditoria Técnica, Operacional e Estratégica  |  23/05/2026  |  Uso interno e confidencial')
r.font.size = Pt(9)
r.font.color.rgb = COR_CINZA
r.italic = True

# ────────────────────────────────────────────────────────────────────────────────
output_path = r'C:\Fluxy\Auditoria_Tecnica_Fluxy_2026.docx'
doc.save(output_path)
print(f'Documento gerado: {output_path}')
