# -*- coding: utf-8 -*-
"""
CADERNO DE ANOTAÇÕES DO TESTE (05/09) — pedido do cliente.

Uma seção por tela, na ORDEM EXATA em que as telas serão abertas, com espaço
pautado para escrever à mão ou digitar. Tudo sai de
scripts/qa-preview/saida/plano-de-teste.json — que vem do manifesto, do diff
do git e da matriz. Nenhum texto aqui é digitado de memória.
"""
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HOJE = '05/09/2026'
AZUL = RGBColor(0x1F, 0x38, 0x64)
CINZA = RGBColor(0x59, 0x59, 0x59)
VERMELHO = RGBColor(0xA6, 0x1C, 0x1C)

plano = json.load(open('scripts/qa-preview/saida/plano-de-teste.json', encoding='utf-8'))
doc = Document()

# ---- página A4, margens de 2cm ------------------------------------------
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
for lado in ('top', 'bottom', 'left', 'right'):
    setattr(sec, f'{lado}_margin', Cm(2))

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)


def par(texto='', tam=10, negrito=False, italico=False, cor=None,
        antes=0, depois=4, recuo=0, alinhamento=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(depois)
    if recuo:
        p.paragraph_format.left_indent = Cm(recuo)
    if alinhamento is not None:
        p.alignment = alinhamento
    r = p.add_run(texto)
    r.font.size = Pt(tam)
    r.bold = negrito
    r.italic = italico
    if cor is not None:
        r.font.color.rgb = cor
    return p


def linha_pautada():
    """Linha para escrever: parágrafo vazio com borda embaixo."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(13)
    pPr = p._p.get_or_add_pPr()
    bordas = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
    b.set(qn('w:space'), '2'); b.set(qn('w:color'), 'AAAAAA')
    bordas.append(b)
    pPr.append(bordas)
    return p


def rodape_com_numero():
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def campo(instr):
        f = OxmlElement('w:fldSimple'); f.set(qn('w:instr'), instr)
        r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = '1'
        r.append(t); f.append(r); return f
    r = p.add_run('Página '); r.font.size = Pt(8); r.font.color.rgb = CINZA
    p._p.append(campo('PAGE'))
    r = p.add_run(' de '); r.font.size = Pt(8); r.font.color.rgb = CINZA
    p._p.append(campo('NUMPAGES'))


cab = sec.header.paragraphs[0]
cab.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rc = cab.add_run('Caderno de teste — reforma do frontend')
rc.font.size = Pt(8); rc.font.color.rgb = CINZA
rodape_com_numero()

# ---- capa ----------------------------------------------------------------
total_telas = sum(len(r['telas']) for r in plano['rodadas'])
par('', depois=60)
par('CADERNO DE TESTE', tam=28, negrito=True, cor=AZUL, depois=2)
par(f'Reforma do frontend — {total_telas} telas em {len(plano["rodadas"])} rodadas',
    tam=14, cor=CINZA, depois=18)
par(f'Gerado em {HOJE} · preview refactor-dev.jrfluxy.com.br', tam=10, cor=CINZA, depois=20)

par('COMO USAR ESTE CADERNO', tam=11, negrito=True, cor=AZUL, antes=14, depois=6)
par('As telas estão na ordem em que serão abertas — começando pelos módulos que você mais usa '
    '(Solicitações, Financeiro, Compras, Cadastros) e terminando nos raros e nas telas de acesso. '
    'Cada seção traz o nome da tela, onde ela fica, o que mudou nela nesta reforma e uma lista do '
    'que analisar ali especificamente.')
par('O que está em "o que analisar" não é genérico: sai do que aquela tela TEM — tabela, coluna de '
    'dinheiro, filtros, janela, abas — e do que a verificação automática NÃO conseguiu provar nela.')
par('O que está em "o que mudou" foi lido do diff do próprio arquivo, não do assunto dos commits: '
    'descreve a tela, não a leva de trabalho.')

par('AS DUAS MARCAS EM VERMELHO', tam=11, negrito=True, cor=AZUL, antes=14, depois=6)
par('SEM DADO — a tela tem a capacidade, mas a base do preview não devolveu registro para '
    'exercitá-la. A verificação automática não conseguiu provar. Se você tiver registro nesta tela, '
    'é o ponto mais importante de olhar.', cor=VERMELHO)
par('CORRIGIDO E NÃO REMEDIDO — foi consertado nesta semana e o conserto ainda não passou por nova '
    'verificação no preview. Vale um olhar extra.', cor=VERMELHO)

par('O QUE ESTE CADERNO NÃO GARANTE', tam=11, negrito=True, cor=AZUL, antes=14, depois=6)
par('A verificação automática mede forma: alinhamento, largura, faixa fixa, contraste, teclado. '
    'Ela não sabe se o número está certo, se a regra do seu negócio foi respeitada ou se a tela faz '
    'sentido para quem trabalha nela. É isso que só você pode dizer — e é para isso que existem as '
    'linhas em branco.')
doc.add_page_break()

# ---- plano de rodadas ----------------------------------------------------
par('Plano de rodadas', tam=16, negrito=True, cor=AZUL, depois=6)
par(f'{total_telas} telas em {len(plano["rodadas"])} rodadas de até 10. '
    'O asterisco (*) marca a tela que tem célula SEM DADO na matriz.', cor=CINZA, depois=8)

tab = doc.add_table(rows=1, cols=5)
tab.style = 'Table Grid'
tab.alignment = WD_TABLE_ALIGNMENT.CENTER
larguras = [Cm(1.3), Cm(3.3), Cm(9.6), Cm(1.1), Cm(1.4)]
titulos = ['Rodada', 'Módulo / tema', 'Telas', 'Qtd', 'S/ dado']
for i, (celula, titulo) in enumerate(zip(tab.rows[0].cells, titulos)):
    celula.width = larguras[i]
    celula.text = ''
    p = celula.paragraphs[0]
    r = p.add_run(titulo); r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), '1F3864')
    celula._tc.get_or_add_tcPr().append(sh)

for idx, r in enumerate(plano['rodadas']):
    com_sd = [t for t in r['telas'] if t['semDado']]
    nomes = ' · '.join((t['nome'] + ' *') if t['semDado'] else t['nome'] for t in r['telas'])
    valores = [str(r['numero']), ' + '.join(r['modulos']), nomes,
               str(len(r['telas'])), str(len(com_sd)) if com_sd else '—']
    linha = tab.add_row()
    for i, (celula, valor) in enumerate(zip(linha.cells, valores)):
        celula.width = larguras[i]
        celula.text = ''
        p = celula.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        run = p.add_run(valor)
        run.font.size = Pt(8.5 if i == 2 else 9)
        if i == 0:
            run.bold = True
        if i in (3, 4):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 4 and com_sd:
            run.bold = True; run.font.color.rgb = VERMELHO
        if idx % 2:
            sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), 'F2F5FA')
            celula._tc.get_or_add_tcPr().append(sh)
doc.add_page_break()

# ---- uma seção por tela --------------------------------------------------
n = 0
for r in plano['rodadas']:
    par(f'Rodada {r["numero"]} — {" + ".join(r["modulos"])}',
        tam=15, negrito=True, cor=AZUL, antes=6, depois=8)
    for t in r['telas']:
        n += 1
        par(f'{n}. {t["nome"]}', tam=13, negrito=True, cor=AZUL, antes=12, depois=3)
        par(f'Onde fica: {t["comoChegar"]["texto"]}', tam=9.5, cor=CINZA, depois=2)

        if t['semDado']:
            par(f'SEM DADO na matriz — {", ".join(t["semDado"]["itens"])}',
                tam=9.5, negrito=True, cor=VERMELHO, depois=2)
        if t['falhou']:
            for f in t['falhou']:
                par(f'CORRIGIDO NESTA SEMANA, ainda não remedido ({f["item"]}): {f["motivo"]}',
                    tam=9.5, negrito=True, cor=VERMELHO, depois=2)

        par('O que mudou nesta reforma', tam=10, negrito=True, antes=6, depois=2)
        for m in t['mudou'][:3]:
            par('— ' + m, tam=9.5, recuo=0.5, depois=1.5)

        par('O que analisar nesta tela', tam=10, negrito=True, antes=6, depois=2)
        for o in t['olhar']:
            par('— ' + o, tam=9.5, recuo=0.5, depois=1.5)

        par('Minhas anotações', tam=10, negrito=True, cor=AZUL, antes=7, depois=7)
        for _ in range(6):
            linha_pautada()
    doc.add_page_break()

doc.save('../docs/CADERNO-DE-TESTE.docx')
print(f'[caderno] {n} seções gravadas em docs/CADERNO-DE-TESTE.docx')
